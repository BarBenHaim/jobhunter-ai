from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading_he(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Arial'
    # shading bar
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'DCE6F1')
    pPr.append(shd)
    return p


def add_bullet_he(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Arial'
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


doc = Document()

# Page margins
for s in doc.sections:
    s.top_margin = Inches(0.6)
    s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

# Title
title = doc.add_paragraph()
set_rtl(title)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run('דף עזר לבוחן בפייתון')
trun.bold = True
trun.font.size = Pt(20)
trun.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
trun.font.name = 'Arial'

# 1
add_heading_he(doc, '1. משתנים וטיפוסי נתונים')
add_bullet_he(doc, 'הטיפוס נקבע אוטומטית לפי הערך.')
add_bullet_he(doc, 'פייתון רגישה לאותיות גדולות/קטנות: x ו-X שונים.')
add_code(doc, '''x = 10                # int
pi = 3.14             # float
name = "Alex"         # str
is_happy = True       # bool (True / False - אות גדולה)

num = int("5")        # המרה למספר שלם
text = str(10)        # המרה למחרוזת''')

# 2
add_heading_he(doc, '2. תנאים (If - Else)')
add_bullet_he(doc, 'הזחה (4 רווחים / Tab) מגדירה את בלוק התנאי.')
add_bullet_he(doc, 'נקודתיים (:) בסוף שורת if / elif / else.')
add_code(doc, '''score = 85

if score >= 90:
    print("A")
elif score >= 80:
    print("B")
else:
    print("F")

# אופרטורים לוגיים: and, or, not
if score > 0 and not is_happy:
    print("Check again")''')

# 3
add_heading_he(doc, '3. רשימות (Lists)')
add_bullet_he(doc, 'אוסף מסודר, האינדקסים מתחילים מ-0.')
add_bullet_he(doc, 'חיתוך [start:end] לא כולל את אינדקס end.')
add_code(doc, '''fruits = ["apple", "banana", "cherry"]

print(fruits[0])      # apple
print(fruits[-1])     # cherry (אחרון)
print(fruits[0:2])    # ['apple', 'banana']

fruits[1] = "blueberry"    # שינוי איבר
fruits.append("orange")    # הוספה לסוף
fruits.insert(1, "mango")  # הוספה באינדקס 1

fruits.remove("apple")     # מחיקה לפי ערך
item = fruits.pop()        # מחיקת האחרון + החזרתו
del fruits[0]              # מחיקה לפי אינדקס''')

# 4
add_heading_he(doc, '4. לולאות (Loops)')
add_bullet_he(doc, 'range(start, stop, step) - רץ עד stop לא כולל.')
add_bullet_he(doc, 'ב-while באחריותנו לקדם את התנאי.')
add_code(doc, '''for fruit in fruits:
    print(fruit)

for i in range(5):           # 0..4
    print(i)

for i in range(1, 10, 2):    # 1,3,5,7,9
    print(i)

for index, fruit in enumerate(fruits):
    print(f"Index {index}: {fruit}")

count = 0
while count < 3:
    print(count)
    count += 1''')

# 5
add_heading_he(doc, '5. פונקציות והחזרת ערך')
add_bullet_he(doc, 'מוגדרת עם def.')
add_bullet_he(doc, 'return מחזיר תוצאה ועוצר מיד את הפונקציה.')
add_code(doc, '''def say_hello(user_name):
    print(f"Hello {user_name}!")

def add_numbers(a, b):
    return a + b

say_hello("Dan")
sum_res = add_numbers(5, 7)   # 12''')

# 6
add_heading_he(doc, '6. פונקציות מובנות ומילונים')
add_bullet_he(doc, 'len(x) - אורך רשימה / מחרוזת.')
add_bullet_he(doc, 'sum(lst) | min/max | sorted(lst) (חדשה) | lst.sort() (במקום).')
add_bullet_he(doc, 'input() קולט מחרוזת. למספר: int(input()).')
add_code(doc, '''student = {"name": "Ron", "age": 23, "grades": [90, 85]}

print(student["name"])        # Ron
student["age"] = 24           # עדכון ערך
student["city"] = "Tel Aviv"  # הוספת מפתח חדש''')

# Summary example
add_heading_he(doc, 'דוגמה מסכמת')
add_code(doc, '''def analyze_grades(grades_list):
    passing_grades = []
    for grade in grades_list:
        if grade >= 60:
            passing_grades.append(grade)

    if len(passing_grades) == 0:
        return {"status": "No one passed", "average": 0}

    avg = sum(passing_grades) / len(passing_grades)
    return {"passed_count": len(passing_grades), "average": avg}

all_grades = [45, 85, 90, 55, 100, 72]
analysis = analyze_grades(all_grades)
print(f"Results: {analysis}")
# Results: {'passed_count': 4, 'average': 86.75}''')

doc.save('/home/user/jobhunter-ai/python_cheat_sheet.docx')
print("saved")
